"""
Export functionality for generating PDF and DOCX reports
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors as reportlab_colors
import io
from datetime import datetime


def export_to_docx(query, answer, sources, agents_output):
    """
    Export conversation to Word document
    
    Args:
        query (str): User's query
        answer (str): AI's answer
        sources (list): List of source dictionaries
        agents_output (dict): Multi-agent output breakdown
        
    Returns:
        BytesIO: Word document as bytes
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('AI Research Co-Pilot Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    doc.add_paragraph('_' * 50)
    
    # Query
    doc.add_heading('Research Query', level=1)
    doc.add_paragraph(query)
    
    # Multi-Agent Breakdown
    if agents_output:
        doc.add_heading('Multi-Agent Analysis', level=1)
        
        if 'research' in agents_output:
            doc.add_heading('1️⃣ Research Agent Findings', level=2)
            doc.add_paragraph(agents_output['research'])
        
        if 'analysis' in agents_output:
            doc.add_heading('2️⃣ Analysis Agent Insights', level=2)
            doc.add_paragraph(agents_output['analysis'])
        
        if 'final' in agents_output:
            doc.add_heading('3️⃣ Final Synthesis', level=2)
            doc.add_paragraph(agents_output['final'])
    else:
        # Single response
        doc.add_heading('AI Response', level=1)
        doc.add_paragraph(answer)
    
    # Sources
    if sources:
        doc.add_heading('Sources & References', level=1)
        for i, source in enumerate(sources, 1):
            p = doc.add_paragraph(f"{i}. ")
            p.add_run(source['title']).bold = True
            doc.add_paragraph(f"   {source['snippet']}")
            doc.add_paragraph(f"   URL: {source['url']}")
    
    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io


def export_to_pdf(query, answer, sources):
    """
    Export conversation to PDF
    
    Args:
        query (str): User's query
        answer (str): AI's answer
        sources (list): List of source dictionaries
        
    Returns:
        BytesIO: PDF document as bytes
    """
    pdf_io = io.BytesIO()
    doc = SimpleDocTemplate(pdf_io, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=reportlab_colors.HexColor('#667eea'),
        spaceAfter=30,
        alignment=1
    )
    story.append(Paragraph("AI Research Co-Pilot Report", title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Metadata
    story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", styles['Normal']))
    story.append(Spacer(1, 0.3*inch))
    
    # Query
    story.append(Paragraph("Research Query", styles['Heading2']))
    story.append(Paragraph(query, styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    
    # Answer
    story.append(Paragraph("AI Response", styles['Heading2']))
    story.append(Paragraph(answer.replace('\n', '<br/>'), styles['Normal']))
    story.append(Spacer(1, 0.3*inch))
    
    # Sources
    if sources:
        story.append(Paragraph("Sources & References", styles['Heading2']))
        for i, source in enumerate(sources, 1):
            story.append(Paragraph(f"{i}. <b>{source['title']}</b>", styles['Normal']))
            story.append(Paragraph(f"{source['snippet']}", styles['Normal']))
            story.append(Paragraph(f"URL: {source['url']}", styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
    
    doc.build(story)
    pdf_io.seek(0)
    
    return pdf_io